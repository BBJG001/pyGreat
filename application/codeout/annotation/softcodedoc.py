from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import re
from application.findfile.getallsub import *


# 返回删掉file文件的的空行和注释的内容
def getcode(file):
    with open(file) as ff:
        ff_ = ff.read()

        ff_ = re.sub('\/\/[^\n]*', '', ff_)  # 替换掉//注释
        ff_ = re.sub('\/\*(\s|.)*?\*\/', '', ff_)  # 替换掉/**/注释

        fflist = ff_.splitlines()  # 按\n分割成列表
        for line in fflist[:]:  # fflist[:]作用是生成一个fflist的副本，解决循环list进行remove()会跳过一些元素的问题
            if line == '':  # 删除掉空行（只含\n的行在splitlines()分割时会被分成空字符''）
                fflist.remove(line)
            if re.match(r'^\s+$', line):  # 删除掉空白字符（如\t）行
                fflist.remove(line)
    res = '\n'.join(fflist) + '\n'
    return res, len(fflist)  # 返回去格式代码和代码行数


# 把flist中所有的文件去注释、去空行保存在savepath中
def saveDocFile(flist, savepath):
    # SINGLE         =>  单倍行距（默认）
    # ONE_POINT_FIVE =>  1.5倍行距
    # DOUBLE2        =>  倍行距
    # AT_LEAST       =>  最小值
    # EXACTLY        =>  固定值
    # MULTIPLE       =>  多倍行距
    doc = Document()
    p = doc.add_paragraph('')  # 增加一页
    doc.styles['Normal'].font.name = 'Times New Roman'  # 正文是normal， 设置正文的字体格式
    doc.styles['Normal'].font.size = Pt(8)  # 设置字体的大小为 5 号字体
    p.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = Pt(12.9)  # 固定值12,9磅, 保证每页有50行代码

    lines = 0
    for i, f in enumerate(flist):
        print('starting deal %d/%d' % (i + 1, len(flist)))
        codef, codelies = getcode(f)
        lines += codelies

        if lines > 3000:    # 如果加上这个文件就超过3000行（就是超过60页了）
            doc.save(savepath)
            return
        else:
            p.add_run(codef)
    doc.save(savepath)      # 不足60 页进行保存
    print('all done')

if __name__ == '__main__':
    path = r'E:\ZYD\document\2017-2018上学期\Web课设\实验室设备管理\sysglxt\src\com'
    savePath = r'E:\code.doc'
    fileList = getAllSub(path)[1]  # 递归获取所有文件
    fileList = getSuffile(fileList, '.java')  # 从中筛选出.java文件
    saveDocFile(fileList, savePath)
