import tkinter as tk
from tkinter.filedialog import *
import os
from docx import Document
from docx.shared import Pt
import re

# fileList = []  # 使用全局列表保存文件路径

def getAllFile(path, fileList):  # 使用递归方法
    dirList = []  # 保存文件夹
    files = os.listdir(path)  # 返回一个列表，其中包含文件 和 文件夹
    for f in files:
        if (os.path.isdir(path + '/' + f)):
            dirList.append(path + '/' + f)  # 将文件夹 名字 进行保存

        if (os.path.isfile(path + '/' + f)):
            fileList.append(path + '/' + f)  # 将文件名 保存

    for dir in dirList:  # 如果文件夹为空时，递归自动退出
        getAllFile(dir, fileList)  # 递归保存到将.java 文件保存到 fileList 中


# getAllFile(r'E:\ZYD\document\2017-2018上学期\Web课设\实验室设备管理\sysglxt\src\com', fileList)
# print('文件数量为： ', len(fileList))


def getJavaFile(fileList):
    for file in fileList:
        if not file.endswith('.java'):  # 删除不是 .java 文件的格式
            fileList.remove(file)
    print('文件数量为： ', len(fileList))


# getJavaFile(fileList)
# print(os.path.isfile(fileList[0]))  # 判断第一个值是否是文件

# 著作权申请表上有 源程序量 一项，需要打印的60页源码是截断输出的，并不是总行数，这里另外做一下
def getAllLines(flist):
    ln = 0
    for ff in flist:
        count = len(open(ff, 'r', errors='ignore').readlines())
        ln+=count
    return ln

# print('源程序量（程序总行数）：', getAllLines(fileList), 'lines')

def saveDocFile(pathin, pathout):
    # SINGLE         =>  单倍行距（默认）
    # ONE_POINT_FIVE =>  1.5倍行距
    # DOUBLE2        =>  倍行距
    # AT_LEAST       =>  最小值
    # EXACTLY        =>  固定值
    # MULTIPLE       =>  多倍行距
    fileList = []  # 使用全局列表保存文件路径

    getAllFile(pathin, fileList)
    getJavaFile(fileList)
    lines = getAllLines(fileList)

    doc = Document()
    from docx.enum.text import WD_LINE_SPACING
    p = doc.add_paragraph('')  # 增加一页
    doc.styles['Normal'].font.name = 'Times New Roman'  # 正文是normal， 设置正文的字体格式
    doc.styles['Normal'].font.size = Pt(8)  # 设置字体的大小为 5 号字体
    p.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = Pt(12.9)  # 固定值12,9磅, 保证每页有50行代码
    save_file = pathout+'\code.doc'
    codeNum = 0
    for i, f in enumerate(fileList):
        print('starting deal %d / %d' % (i+1, len(fileList)))
        with open(f, encoding='UTF-8', errors='ignore') as file:  # 转换编码以实现正确输出中文格式
            for line in file.readlines():

                if line == '\n':  # 删除空行
                    continue
                if re.match(r'^\s+$', line):  # 使用正则表达式删除全是空格的空行
                    continue
                if line.__contains__(r'/*') or \
                        line.__contains__(r' *'):  # 删除注释
                    continue
                if line.__contains__(r'//'):  # 删除包含 // 的注释, 严格意义上应该使用正则表达式进行删除
                    continue
                p.add_run(line)
                codeNum += 1  # 记录是已经写入的数据
                if codeNum == 3050:  # 保证打印出不大大超过与 60 页
                    doc.save(save_file)
                    return
    doc.save(save_file)  # 不足60 页进行保存
    return lines

# saveDocFile()
# print('all done')

if __name__ == '__main__':

    # 点击“选择文件夹”调用该功能
    def selectDirecPathIn():
        path_ = askdirectory(title='选择文件夹')
        pathin.set(path_)

    # 点击“选择文件夹”调用该功能
    def selectDirecPathOut():
        path_ = askdirectory(title='选择文件夹')
        pathout.set(path_)


    # 点击执行调用该功能
    def executeit():
        hint.set('Processing . . .')
        ln = saveDocFile(pathin.get(), pathout.get())
        hint.set('convert finished\n程序行数：'+str(ln))

    # 生成窗口
    window = tk.Tk()
    window.title('软著代码整理器')
    window.geometry('450x300')


    # 格式转换label
    tk.Label(window, text='软著代码整理器', height=2).pack(anchor=CENTER)  # x是从左向右的偏移，y是从上向下的偏移

    tk.Label(window, width=50, text='说明：该程序可以递归获得输入文件夹下的所有java文件，将其中\n代码去注释、去空行后生成code.doc文件保存到指定输出文件夹下').place(x=50, y=40)  # 从左偏，从上偏

    # 输入文件一行
    pathin = tk.StringVar()  # 定义变量
    tk.Label(window, text='输入文件:').place(x=50, y=100)
    entry_pathin = tk.Entry(window, textvariable=pathin).place(x=160, y=100)  # 输入框
    btn_pathin = tk.Button(window, text='选择文件夹', command=selectDirecPathIn).place(x=340, y=95)  # 按钮

    # 输出文件夹的一行
    pathout = tk.StringVar()
    tk.Label(window, text='输出位置:').place(x=50, y=150)  # 从左偏，从上偏
    tk.Entry(window, textvariable=pathout).place(x=160, y=150)
    tk.Button(window, text='选择文件夹', command=selectDirecPathOut).place(x=335, y=145)
    tk.Label(window, text='默认在指定文件夹下生成code.doc').place(x=160, y=180)  # 从左偏，从上偏

    # 执行按钮
    tk.Button(window, text='  执  行  ', command=executeit, width=10).place(x=190, y=215)

    # 提示区域
    hint = tk.StringVar()
    hint.set('')
    tk.Label(window, textvariable=hint).place(x=190, y=250)

    # 不停的刷新显示
    window.mainloop()