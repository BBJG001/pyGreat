#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
__title__ = ''
__author__ = 'mike_jun'
__mtime__ = '2019-7-1'
#目的： 1. 先将一个文件夹下的所有文件夹的 .java 文件路径保存到一个列表中
        2. 依次读取列表的路径， 将 .java 文件内容保存到word 中
"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import re
from docx.shared import Length

fileList = []  # 使用全局列表保存文件路径


def getAllFile(path, fileList):  # 使用递归方法
    dirList = []  # 保存文件夹
    files = os.listdir(path)  # 返回一个列表，其中包含文件 和 文件夹
    for f in files:
        if (os.path.isdir(path + '/' + f)):
            dirList.append(path + '/' + f)  # 将文件夹 名字 进行保存

        if (os.path.isfile(path + '/' + f)):
            fileList.append(path + '/' + f)  # 将文件名 保存

    for dir in dirList:  # 如果文件夹为空时，递归自动退出
        getAllFile(dir, fileList)  # 递归保存到将.java 文件保存到 fileList 中


getAllFile(r'E:\ZYD\document\2017-2018上学期\Web课设\实验室设备管理\sysglxt\src\com', fileList)
print('文件数量为： ', len(fileList))


def getJavaFile(fileList):
    for file in fileList[:]:
        if not file.endswith('.java'):  # 删除不是 .java 文件的格式
            fileList.remove(file)
    print('文件数量为： ', len(fileList))


getJavaFile(fileList)
print(os.path.isfile(fileList[0]))  # 判断第一个值是否是文件

# 著作权申请表上有 源程序量 一项，需要打印的60页源码是截断输出的，并不是总行数，这里另外做一下
def getAllLines(flist):
    ln = 0
    for ff in flist:
        count = len(open(ff, 'r').readlines())
        ln+=count
    return ln

print('源程序量（程序总行数）：', getAllLines(fileList), 'lines')

def saveDocFile():
    # SINGLE         =>  单倍行距（默认）
    # ONE_POINT_FIVE =>  1.5倍行距
    # DOUBLE2        =>  倍行距
    # AT_LEAST       =>  最小值
    # EXACTLY        =>  固定值
    # MULTIPLE       =>  多倍行距
    doc = Document()
    from docx.enum.text import WD_LINE_SPACING
    p = doc.add_paragraph('')  # 增加一页
    doc.styles['Normal'].font.name = 'Times New Roman'  # 正文是normal， 设置正文的字体格式
    doc.styles['Normal'].font.size = Pt(8)  # 设置字体的大小为 5 号字体
    p.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 固定值
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = Pt(12.9)  # 固定值12,9磅, 保证每页有50行代码
    save_file = r'E:\code.doc'
    codeNum = 0
    for i, f in enumerate(fileList):
        print('starting deal %d / %d' % (i+1, len(fileList)))
        with open(f, encoding='UTF-8', errors='ignore') as file:  # 转换编码以实现正确输出中文格式
            for line in file.readlines():

                if line == '\n':  # 删除空行
                    continue
                if re.match(r'^\s+$', line):  # 使用正则表达式删除全是空格的空行
                    continue
                if line.__contains__(r'/*') or \
                        line.__contains__(r' *'):  # 删除注释
                    continue
                if line.__contains__(r'//'):  # 删除包含 // 的注释, 严格意义上应该使用正则表达式进行删除
                    continue
                p.add_run(line)
                codeNum += 1  # 记录是已经写入的数据
                if codeNum == 3050:  # 保证打印出不大大超过与 60 页
                    doc.save(save_file)
                    return
    doc.save(save_file)  # 不足60 页进行保存
    print('all done')

saveDocFile()
print('all done')